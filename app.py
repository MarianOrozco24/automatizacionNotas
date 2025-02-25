import customtkinter as ctk
import pandas as pd
from tkinter import messagebox as mx
from procesamiento_notas.procesador import limpiar_df, analisis_df
from scrapeo import scrapeo
from xls_a_xlsx.transformation import progress
from xls_a_xlsx.transformation import buscar_archivos
from config import Config
from docx import Document
import os

def main():
    # Descarga de informacion
    scrapeo()
    try:
        archivos_descargados = buscar_archivos(Config.ruta_directorio_input)
        for archivo in archivos_descargados:
            # Luego de terminado el proceso de tranformacion de xls a xlsx extraemos el nombre resultado
            nombre_excel = progress(archivo)
            '''' ---------- Esto lo uso para hacer pruebas sin scrapear -----------'''
            # nomnre_excel = os.listdir(Config.ruta_directorio_input)
            # nombre_excel = nomnre_excel[0]
            # # Lo convertimos en df
            ''' ----------------------------------------------------------------- '''

            df = pd.read_excel(f"{Config.ruta_directorio_input}/{nombre_excel}")

            # Lo limpiamos
            df_limpio = limpiar_df(df)

            # Exportamos para comprobar
            resultado = analisis_df(df_limpio)

            # Renombramos el excel para que no se repita el xlsx
            nombre_excel = nombre_excel.replace(".xlsx", "")


            # Guardar resultado en Excel
            resultado.to_excel(f"{Config.ruta_salida}/xlsx/{nombre_excel}.xlsx", engine="openpyxl", index=False)
            


            try:
                nombre_excel = nombre_excel.replace(".xlsx", "")
                # Seleccionar solo las dos columnas necesarias
                df_a_word = resultado[["Nombre y Apellido", "Cantidad materias"]]

                # Crear el documento Word
                nombre_word = nombre_excel.replace(".xlsx", ".docx")
                df_word = Document()
                df_word.add_heading(nombre_excel, level=1)

                # Crear la tabla con una fila para los encabezados
                tabla = df_word.add_table(rows=1, cols=len(df_a_word.columns))

                # Agregar encabezados
                encabezados = tabla.rows[0].cells
                for i, columna in enumerate(df_a_word.columns):
                    encabezados[i].text = columna

                # Agregar los datos del DataFrame a la tabla
                for _, fila in df_a_word.iterrows():
                    row_cells = tabla.add_row().cells
                    for i, valor in enumerate(fila):
                        row_cells[i].text = str(valor)

                # Guardar el documento
                df_word.save(f"{Config.ruta_salida}/docx/{nombre_word}.docx")
                print(f"✅ Archivo Word '{nombre_word}' generado con éxito.")
            except Exception as e:
                print("❌ ERROR: al intentar exporta a word ", e)
            
        try:
            # Remover archivos de input
            archivos_xlsx_input = os.listdir(Config.ruta_directorio_input)
            for archivo in archivos_xlsx_input:
                os.remove(f"{Config.ruta_directorio_input}/{nombre_excel}.xlsx")
                if os.remove:
                    print("🗑 Se removio con exito el elemento de INPUTS")

        except Exception as e:
            print("❌ ERROR: Se produjo un error al intentar eliminar el elemento de INPUTS. Msj Error: ", e)

                        
        
        mx.showinfo("Filtrado completo", "Se han exportado todos los archivos a la carpeta outputs")
            
    except Exception as e:
        print("❌ ERROR: Al exportar el excel", e)

# interfaz grafica
ventana_principal = ctk.CTk()
ventana_principal.title("Analisis Notas")
ventana_principal.geometry("850x450")

# label_title
label_title = ctk.CTkLabel(ventana_principal, text="Analisis Notas", font=("Helvetica", 20, "bold"))
label_title.pack(pady=35, padx=5)


# Botton continuar 
frame_continuar = ctk.CTkFrame(ventana_principal, fg_color="#00ADB5")
frame_continuar.pack(pady=5, padx=5)
button_continue = ctk.CTkButton(frame_continuar, text="Procesar", fg_color="#2E2E2E", command=lambda:main())
button_continue.pack(pady=2, padx=2)

ventana_principal.mainloop()






